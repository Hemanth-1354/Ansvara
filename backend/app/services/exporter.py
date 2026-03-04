import io
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors


def export_to_docx(run_data: Dict, answers: List[Dict]) -> bytes:
    """Generate a Word document with all Q&A pairs."""
    doc = Document()

    # Title
    title = doc.add_heading("Questionnaire Answers Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    summary = run_data.get("summary", {})
    if summary:
        doc.add_paragraph(
            f"Summary: {summary.get('total', 0)} questions | "
            f"{summary.get('answered', 0)} answered | "
            f"{summary.get('not_found', 0)} not found"
        )

    doc.add_paragraph("")

    for i, ans in enumerate(answers, 1):
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. {ans['question_text']}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run("Answer: ")
        a_run.bold = True
        a_para.add_run(ans.get("answer_text") or "Not found in references.")

        # Confidence
        if ans.get("confidence") is not None and ans.get("is_found"):
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(f"Confidence: {int(ans['confidence'] * 100)}%")
            conf_run.italic = True
            conf_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Citations
        citations = ans.get("citations") or []
        if citations:
            cit_para = doc.add_paragraph()
            cit_run = cit_para.add_run("Citations: ")
            cit_run.bold = True
            cit_run.italic = True
            for cit in citations:
                doc.add_paragraph(
                    f"  • [{cit['doc_name']}] — {cit['snippet']}",
                    style="List Bullet"
                )

        doc.add_paragraph("")
        if i < len(answers):
            doc.add_paragraph("─" * 60)
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_pdf(run_data: Dict, answers: List[Dict]) -> bytes:
    """Generate a PDF document with all Q&A pairs."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Title'], spaceAfter=12)
    q_style = ParagraphStyle('question', parent=styles['Normal'],
                              fontName='Helvetica-Bold', fontSize=11, spaceAfter=4,
                              textColor=colors.HexColor('#1a1a2e'))
    a_style = ParagraphStyle('answer', parent=styles['Normal'],
                              fontSize=10, spaceAfter=4, leftIndent=12)
    cit_style = ParagraphStyle('citation', parent=styles['Normal'],
                                fontSize=9, textColor=colors.HexColor('#555555'),
                                leftIndent=24, spaceAfter=2)
    conf_style = ParagraphStyle('confidence', parent=styles['Normal'],
                                 fontSize=9, textColor=colors.HexColor('#2ecc71'),
                                 leftIndent=12, spaceAfter=4, fontName='Helvetica-Oblique')

    story = []
    story.append(Paragraph("Questionnaire Answers Report", title_style))

    summary = run_data.get("summary", {})
    if summary:
        story.append(Paragraph(
            f"<b>Total:</b> {summary.get('total', 0)} | "
            f"<b>Answered:</b> {summary.get('answered', 0)} | "
            f"<b>Not Found:</b> {summary.get('not_found', 0)}",
            styles['Normal']
        ))
    story.append(Spacer(1, 12))

    for i, ans in enumerate(answers, 1):
        story.append(Paragraph(f"Q{i}. {ans['question_text']}", q_style))

        answer_text = ans.get("answer_text") or "Not found in references."
        story.append(Paragraph(f"<b>Answer:</b> {answer_text}", a_style))

        if ans.get("confidence") is not None and ans.get("is_found"):
            story.append(Paragraph(
                f"Confidence: {int(ans['confidence'] * 100)}%", conf_style
            ))

        citations = ans.get("citations") or []
        if citations:
            story.append(Paragraph("<i><b>Citations:</b></i>", cit_style))
            for cit in citations:
                story.append(Paragraph(
                    f"• [{cit['doc_name']}] — {cit['snippet'][:120]}...", cit_style
                ))

        story.append(Spacer(1, 8))
        if i < len(answers):
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()
